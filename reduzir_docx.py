#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
reduzir_docx.py — Reduz o tamanho de relatórios .docx recomprimindo as fotos
embutidas em word/media, SEM alterar layout, texto, legendas ou posição das
imagens.

Como funciona
-------------
Um .docx é um arquivo ZIP. O tamanho de EXIBIÇÃO de cada imagem é definido no
XML do documento (word/document.xml etc.) e NÃO muda quando o arquivo de
imagem em word/media é substituído por outro de menor resolução com o MESMO
nome. Este script explora exatamente essa propriedade:

  * Nenhum XML do documento é tocado — os bytes de todas as partes que não
    são imagem são copiados intactos (verificado por CRC na validação).
  * Apenas os arquivos JPEG/PNG de word/media são redimensionados (máx.
    1600 px na maior dimensão) e recomprimidos (JPEG q=80; PNG otimizado),
    mantendo exatamente o mesmo nome e formato (sem conversão PNG->JPEG).
  * EXIF (orientação) e perfil ICC dos JPEGs são preservados, para que a
    foto continue sendo exibida exatamente como antes.

Regras de segurança (REGRA DE OURO: integridade acima de redução):
  * O original NUNCA é alterado (aberto somente para leitura).
  * Imagens vetoriais (EMF/WMF) e imagens < 200 KB são ignoradas.
  * Se a versão recomprimida ficar MAIOR que a original, mantém-se a original.
  * Todo arquivo gerado passa por validação obrigatória (python-docx abre sem
    erros; mesmo nº de parágrafos, tabelas e imagens; partes não-imagem
    byte-a-byte idênticas via CRC). Se qualquer verificação falhar, o gerado
    é descartado e uma cópia INTACTA do original vai para a pasta de saída.

Uso (Windows):
    py reduzir_docx.py
    py reduzir_docx.py --entrada "C:\\Temp\\RNCs\\Originais" --saida "C:\\Temp\\RNCs\\Compactados"

Requisitos: Python 3.8+, Pillow, python-docx  (py -m pip install pillow python-docx)
"""

import argparse
import gc
import io
import os
import shutil
import sys
import traceback
import zipfile
from datetime import datetime

try:
    from PIL import Image
except ImportError:
    sys.exit("ERRO: Pillow não instalado. Execute:  py -m pip install pillow")
try:
    import docx  # python-docx
except ImportError:
    sys.exit("ERRO: python-docx não instalado. Execute:  py -m pip install python-docx")

# ----------------------------- Configuração ---------------------------------

PASTA_ENTRADA_PADRAO = r"C:\Temp\RNCs\Originais"
PASTA_SAIDA_PADRAO = r"C:\Temp\RNCs\Compactados"

DIM_MAXIMA = 1600            # maior dimensão (px) após redimensionamento
QUALIDADE_JPEG = 80          # qualidade de recompressão JPEG
TAMANHO_MINIMO = 200 * 1024  # imagens menores que isso são ignoradas (200 KB)
LIMITE_ALERTA = 100 * 1024 * 1024  # 100 MB — limiar do resumo final
_CHUNK = 4 * 1024 * 1024     # cópia em blocos de 4 MB (economia de memória)

# Bomba de descompressão / sanidade: não tenta abrir imagens absurdas
Image.MAX_IMAGE_PIXELS = 300_000_000

EXT_JPEG = {".jpg", ".jpeg", ".jpe", ".jfif"}
EXT_PNG = {".png"}
EXT_VETORIAL = {".emf", ".wmf", ".svg", ".wdp"}

# ------------------------------- Utilitários ---------------------------------


def _mb(n_bytes):
    return n_bytes / (1024 * 1024)


def _fmt_mb(n_bytes):
    return f"{_mb(n_bytes):,.1f} MB".replace(",", "X").replace(".", ",").replace("X", ".")


def _eh_media(nome):
    return nome.replace("\\", "/").lower().startswith("word/media/")


def _ext(nome):
    return os.path.splitext(nome)[1].lower()


# --------------------------- Compressão de imagem -----------------------------


def recomprimir_imagem(dados, nome):
    """Tenta recomprimir uma imagem JPEG/PNG.

    Retorna (novos_bytes, None) em caso de sucesso ou (None, motivo) quando a
    imagem deve ser mantida como está.
    """
    ext = _ext(nome)
    with Image.open(io.BytesIO(dados)) as img:
        formato = img.format  # formato REAL do conteúdo, não da extensão
        if formato not in ("JPEG", "PNG"):
            return None, f"formato {formato or 'desconhecido'} (mantida)"
        # Nunca converter formatos nesta passada: conteúdo e extensão precisam
        # ser coerentes (um PNG chamado image1.jpeg seria arriscado recomprimir).
        if formato == "JPEG" and ext not in EXT_JPEG:
            return None, "extensão não confere com o conteúdo (mantida)"
        if formato == "PNG" and ext not in EXT_PNG:
            return None, "extensão não confere com o conteúdo (mantida)"
        if getattr(img, "is_animated", False):
            return None, "imagem animada (mantida)"

        img.load()
        largura, altura = img.size
        maior = max(largura, altura)
        if maior > DIM_MAXIMA:
            escala = DIM_MAXIMA / maior
            novo_tam = (max(1, round(largura * escala)), max(1, round(altura * escala)))
            # PNG com paleta: converter antes de redimensionar para não degradar
            if img.mode == "P":
                img = img.convert("RGBA" if "transparency" in img.info else "RGB")
            img = img.resize(novo_tam, Image.LANCZOS)

        buf = io.BytesIO()
        if formato == "JPEG":
            parametros = {"format": "JPEG", "quality": QUALIDADE_JPEG, "optimize": True}
            # Preserva orientação (EXIF) e cores (ICC) para exibição idêntica
            exif = img.info.get("exif")
            if exif:
                parametros["exif"] = exif
            icc = img.info.get("icc_profile")
            if icc:
                parametros["icc_profile"] = icc
            if img.mode not in ("RGB", "L", "CMYK", "YCbCr"):
                img = img.convert("RGB")
            img.save(buf, **parametros)
        else:  # PNG
            parametros = {"format": "PNG", "optimize": True}
            icc = img.info.get("icc_profile")
            if icc:
                parametros["icc_profile"] = icc
            img.save(buf, **parametros)
        return buf.getvalue(), None


# ------------------------------ Processamento --------------------------------


def _copiar_zipinfo(item):
    novo = zipfile.ZipInfo(filename=item.filename, date_time=item.date_time)
    novo.compress_type = zipfile.ZIP_DEFLATED
    novo.external_attr = item.external_attr
    novo.create_system = item.create_system
    return novo


def processar_docx(origem, destino_tmp, log):
    """Gera destino_tmp a partir de origem, recomprimindo imagens de word/media.

    Retorna um dicionário de estatísticas. Lê a origem SOMENTE para leitura.
    """
    stats = {"processadas": 0, "puladas": 0, "detalhes": []}
    with zipfile.ZipFile(origem, "r") as zin, \
            zipfile.ZipFile(destino_tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            nome = item.filename
            if item.is_dir():
                zout.writestr(_copiar_zipinfo(item), b"")
                continue

            if _eh_media(nome):
                ext = _ext(nome)
                if ext in EXT_VETORIAL:
                    stats["puladas"] += 1
                    stats["detalhes"].append(f"    - {nome}: vetorial (mantida)")
                elif ext not in (EXT_JPEG | EXT_PNG):
                    stats["puladas"] += 1
                    stats["detalhes"].append(f"    - {nome}: extensão {ext or '?'} (mantida)")
                elif item.file_size < TAMANHO_MINIMO:
                    stats["puladas"] += 1
                    stats["detalhes"].append(
                        f"    - {nome}: {item.file_size // 1024} KB < 200 KB (mantida)")
                else:
                    dados = zin.read(nome)
                    try:
                        novos, motivo = recomprimir_imagem(dados, nome)
                    except Exception as exc:  # imagem problemática: manter original
                        novos, motivo = None, f"erro ao processar ({exc}) — mantida"
                    if novos is not None and len(novos) < len(dados):
                        zout.writestr(_copiar_zipinfo(item), novos)
                        stats["processadas"] += 1
                        stats["detalhes"].append(
                            f"    - {nome}: {_fmt_mb(len(dados))} -> {_fmt_mb(len(novos))}")
                        del dados, novos
                        continue
                    stats["puladas"] += 1
                    if novos is not None:
                        motivo = "recompressão ficou maior (original mantida)"
                    stats["detalhes"].append(f"    - {nome}: {motivo}")
                    zout.writestr(_copiar_zipinfo(item), dados)
                    del dados
                    continue
                # media mantida sem alteração (vetorial/pequena/outra extensão)
                with zin.open(item) as fin, zout.open(_copiar_zipinfo(item), "w") as fout:
                    shutil.copyfileobj(fin, fout, _CHUNK)
            else:
                # Partes não-imagem: cópia byte a byte, em blocos (sem carregar
                # o arquivo inteiro na memória)
                with zin.open(item) as fin, zout.open(_copiar_zipinfo(item), "w") as fout:
                    shutil.copyfileobj(fin, fout, _CHUNK)
    return stats


# -------------------------------- Validação ----------------------------------


def _contagens_docx(caminho):
    documento = docx.Document(caminho)
    contagens = (
        len(documento.paragraphs),
        len(documento.tables),
        len(documento.inline_shapes),
    )
    del documento
    gc.collect()
    return contagens


def validar(origem, gerado):
    """Validação obrigatória. Retorna (ok, lista_de_mensagens)."""
    mensagens = []

    # 1) Estrutura do ZIP: mesmos membros; partes não-imagem idênticas (CRC)
    with zipfile.ZipFile(origem, "r") as z1, zipfile.ZipFile(gerado, "r") as z2:
        if z2.testzip() is not None:
            return False, ["ZIP gerado corrompido (testzip falhou)"]
        info1 = {i.filename: i for i in z1.infolist()}
        info2 = {i.filename: i for i in z2.infolist()}
        if set(info1) != set(info2):
            return False, ["lista de membros do ZIP difere do original"]
        n_imagens1 = sum(1 for n in info1 if _eh_media(n))
        n_imagens2 = sum(1 for n in info2 if _eh_media(n))
        if n_imagens1 != n_imagens2:
            return False, [f"nº de imagens difere: {n_imagens1} vs {n_imagens2}"]
        for nome, i1 in info1.items():
            if not _eh_media(nome):
                i2 = info2[nome]
                if i1.CRC != i2.CRC or i1.file_size != i2.file_size:
                    return False, [f"parte não-imagem alterada: {nome}"]
        mensagens.append(f"imagens: {n_imagens1} = {n_imagens2} OK; partes XML intactas (CRC) OK")

    # 2) python-docx abre sem erros e contagens idênticas
    paragrafos1, tabelas1, formas1 = _contagens_docx(origem)
    paragrafos2, tabelas2, formas2 = _contagens_docx(gerado)
    if (paragrafos1, tabelas1, formas1) != (paragrafos2, tabelas2, formas2):
        return False, [
            "contagens diferem: parágrafos {} vs {}, tabelas {} vs {}, imagens inline {} vs {}".format(
                paragrafos1, paragrafos2, tabelas1, tabelas2, formas1, formas2)
        ]
    mensagens.append(
        f"parágrafos: {paragrafos1} OK; tabelas: {tabelas1} OK; imagens inline: {formas1} OK")
    return True, mensagens


# --------------------------------- Pipeline ----------------------------------


def executar(pasta_entrada, pasta_saida):
    if not os.path.isdir(pasta_entrada):
        sys.exit(f"ERRO: pasta de entrada não existe: {pasta_entrada}")
    os.makedirs(pasta_saida, exist_ok=True)

    arquivos = sorted(
        f for f in os.listdir(pasta_entrada)
        if f.lower().endswith(".docx") and not f.startswith("~$")
    )
    if not arquivos:
        sys.exit(f"Nenhum .docx encontrado em {pasta_entrada}")

    caminho_log = os.path.join(pasta_saida, "log_reducao.txt")
    linhas_log = [
        "LOG DE REDUÇÃO DE .DOCX — " + datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
        f"Entrada: {pasta_entrada}",
        f"Saída:   {pasta_saida}",
        f"Parâmetros: máx. {DIM_MAXIMA} px | JPEG q={QUALIDADE_JPEG} | mínimo {TAMANHO_MINIMO // 1024} KB",
        "=" * 78,
    ]
    resultados = []

    for indice, nome_arquivo in enumerate(arquivos, 1):
        origem = os.path.join(pasta_entrada, nome_arquivo)
        destino = os.path.join(pasta_saida, nome_arquivo)
        destino_tmp = destino + ".tmp"
        tamanho_antes = os.path.getsize(origem)
        print(f"[{indice}/{len(arquivos)}] {nome_arquivo} ({_fmt_mb(tamanho_antes)}) ...",
              flush=True)
        linhas_log.append(f"\nARQUIVO: {nome_arquivo}")
        linhas_log.append(f"  Tamanho antes:  {_fmt_mb(tamanho_antes)}")

        falha = None
        stats = {"processadas": 0, "puladas": 0, "detalhes": []}
        try:
            stats = processar_docx(origem, destino_tmp, linhas_log)
            ok, msgs_validacao = validar(origem, destino_tmp)
            if ok:
                os.replace(destino_tmp, destino)
                for msg in msgs_validacao:
                    linhas_log.append(f"  Validação: {msg}")
            else:
                falha = "; ".join(msgs_validacao)
        except Exception:
            falha = "exceção durante o processamento:\n" + traceback.format_exc()

        if falha is not None:
            # Descarta o gerado e mantém cópia INTACTA do original na saída
            if os.path.exists(destino_tmp):
                os.remove(destino_tmp)
            shutil.copy2(origem, destino)
            linhas_log.append(f"  *** FALHA: {falha}")
            linhas_log.append("  *** Ação: arquivo gerado descartado; cópia intacta do "
                              "original mantida na pasta de saída.")
            print(f"    FALHOU — original copiado intacto. Motivo: {falha.splitlines()[0]}")

        tamanho_depois = os.path.getsize(destino)
        reducao = (1 - tamanho_depois / tamanho_antes) * 100 if tamanho_antes else 0.0
        linhas_log.append(f"  Tamanho depois: {_fmt_mb(tamanho_depois)}")
        linhas_log.append(f"  Redução:        {reducao:.1f} %")
        linhas_log.append(f"  Imagens processadas: {stats['processadas']} | puladas: {stats['puladas']}")
        linhas_log.extend(stats["detalhes"])
        print(f"    {_fmt_mb(tamanho_antes)} -> {_fmt_mb(tamanho_depois)} "
              f"({reducao:.1f} % de redução) | imagens: {stats['processadas']} processadas, "
              f"{stats['puladas']} puladas")

        resultados.append({
            "nome": nome_arquivo,
            "antes": tamanho_antes,
            "depois": tamanho_depois,
            "processadas": stats["processadas"],
            "puladas": stats["puladas"],
            "falha": falha,
        })
        gc.collect()  # libera memória antes do próximo arquivo (arquivos de 200+ MB)

    # ------------------------------ Resumo geral -----------------------------
    total_antes = sum(r["antes"] for r in resultados)
    total_depois = sum(r["depois"] for r in resultados)
    falhas = [r for r in resultados if r["falha"]]
    grandes = [r for r in resultados if r["depois"] > LIMITE_ALERTA]

    resumo = [
        "",
        "=" * 78,
        "RESUMO GERAL",
        f"  Arquivos processados: {len(resultados)} | com falha (original mantido): {len(falhas)}",
        f"  Total antes:  {_fmt_mb(total_antes)}",
        f"  Total depois: {_fmt_mb(total_depois)}",
        f"  Redução total: {(1 - total_depois / total_antes) * 100 if total_antes else 0:.1f} %",
    ]
    if grandes:
        resumo.append(f"\n  ATENÇÃO: {len(grandes)} arquivo(s) ainda acima de 100 MB:")
        for r in grandes:
            caminho = os.path.join(pasta_saida, r["nome"])
            peso_png, n_png = _peso_pngs(caminho)
            percentual = (peso_png / r["depois"] * 100) if r["depois"] else 0
            resumo.append(
                f"    - {r['nome']}: {_fmt_mb(r['depois'])} | PNGs em word/media: "
                f"{n_png} arquivo(s), {_fmt_mb(peso_png)} ({percentual:.0f} % do peso)")
        resumo.append(
            "  -> Se autorizado, uma 2ª passada pode converter PNGs fotográficos em JPEG\n"
            "     (com atualização das referências internas e a mesma validação).")
    else:
        resumo.append("\n  Nenhum arquivo restou acima de 100 MB.")

    linhas_log.extend(resumo)
    with open(caminho_log, "w", encoding="utf-8") as arquivo_log:
        arquivo_log.write("\n".join(linhas_log) + "\n")
    print("\n".join(resumo))
    print(f"\nLog completo: {caminho_log}")
    return resultados


def _peso_pngs(caminho_docx):
    """Peso (bytes comprimidos dentro do ZIP) dos PNGs de word/media."""
    peso, quantidade = 0, 0
    with zipfile.ZipFile(caminho_docx, "r") as z:
        for item in z.infolist():
            if _eh_media(item.filename) and _ext(item.filename) in EXT_PNG:
                peso += item.compress_size
                quantidade += 1
    return peso, quantidade


def main():
    parser = argparse.ArgumentParser(
        description="Reduz .docx recomprimindo fotos de word/media sem alterar o documento.")
    parser.add_argument("--entrada", default=PASTA_ENTRADA_PADRAO,
                        help=f"pasta com os .docx originais (padrão: {PASTA_ENTRADA_PADRAO})")
    parser.add_argument("--saida", default=PASTA_SAIDA_PADRAO,
                        help=f"pasta de saída (padrão: {PASTA_SAIDA_PADRAO})")
    argumentos = parser.parse_args()
    executar(argumentos.entrada, argumentos.saida)


if __name__ == "__main__":
    main()
