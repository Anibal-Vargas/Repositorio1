"""Etapas 3 e 4 — Laudos (.docx) preenchendo os modelos do cliente.

Estratégia: usar os modelos ``Laudo_Geral_Padrao.docx`` e
``Laudo_Individual_Padrao.docx`` como base (preservando figuras, metodologia,
normas e formatação) e substituir apenas os campos variáveis. A substituição é
feita de forma robusta, mesmo quando o texto está dividido em vários "runs"
do Word, preservando a formatação ao redor.

Os dados fixos vêm da :class:`Configuracao` (preenchida uma vez); a data das
medições e os dados por máquina vêm da inspeção.
"""

from __future__ import annotations

import io
import os

import docx

from .configuracao import Configuracao
from .imagens import encaixar
from .modelo import FORA_DE_FAIXA, Equipamento, Pacote
from .recursos import caminho_recurso

MODELO_GERAL = caminho_recurso("modelos", "Laudo_Geral_Padrao.docx")
MODELO_INDIVIDUAL = caminho_recurso("modelos", "Laudo_Individual_Padrao.docx")

_MESES = [
    "janeiro", "fevereiro", "março", "abril", "maio", "junho",
    "julho", "agosto", "setembro", "outubro", "novembro", "dezembro",
]


def data_por_extenso(dt) -> str:
    return f"{dt.day:02d} de {_MESES[dt.month - 1]} de {dt.year}"


def _fmt(valor: float) -> str:
    """Formata número em pt-BR com uma casa decimal (57.4 -> '57,4')."""
    return f"{valor:.1f}".replace(".", ",")


# --- Substituição robusta de texto (atravessa runs) ----------------------

def _substituir_no_paragrafo(paragrafo, antigo: str, novo: str) -> bool:
    """Substitui a primeira ocorrência de ``antigo`` por ``novo`` no parágrafo,
    mesmo que o texto esteja dividido entre vários runs. Preserva a formatação
    do run onde a correspondência começa. Devolve True se substituiu."""
    runs = paragrafo.runs
    if not runs:
        return False
    texto = "".join(r.text for r in runs)
    inicio = texto.find(antigo)
    if inicio == -1:
        return False
    fim = inicio + len(antigo)

    pos = 0
    inserido = False
    for r in runs:
        ini_r, fim_r = pos, pos + len(r.text)
        pos = fim_r
        if fim_r <= inicio or ini_r >= fim:
            continue  # run sem sobreposição com o trecho
        ls = max(inicio, ini_r) - ini_r
        le = min(fim, fim_r) - ini_r
        antes, depois = r.text[:ls], r.text[le:]
        if not inserido:
            r.text = antes + novo + depois
            inserido = True
        else:
            r.text = antes + depois
    return True


def _paragrafos_de(container):
    """Todos os parágrafos de um container, inclusive dentro de tabelas."""
    yield from container.paragraphs
    for tabela in container.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                yield from _paragrafos_de(celula)


def _iter_paragrafos(doc):
    yield from _paragrafos_de(doc)
    for secao in doc.sections:
        # Cabeçalhos/rodapés também podem ter tabelas (é onde fica a proposta).
        for cab in (secao.header, secao.footer,
                    secao.first_page_header, secao.first_page_footer,
                    secao.even_page_header, secao.even_page_footer):
            if cab is not None:
                yield from _paragrafos_de(cab)


def _aplicar_substituicoes(doc, substituicoes: dict) -> None:
    # Ordena por comprimento decrescente para evitar substituir trechos menores
    # antes dos maiores que os contêm.
    itens = sorted(substituicoes.items(), key=lambda kv: -len(kv[0]))
    for paragrafo in _iter_paragrafos(doc):
        for antigo, novo in itens:
            if antigo and antigo != novo:
                _substituir_no_paragrafo(paragrafo, antigo, str(novo))


def _ohm_do_documento(doc, padrao: str = "Ω") -> str:
    """Descobre o caractere de ohm usado no documento (U+2126 ou U+03A9)."""
    for p in doc.paragraphs:
        for ch in p.text:
            if ch in ("Ω", "Ω"):
                return ch
    return padrao


# Frase usada no laudo geral quando há mais de um valor de prolongador.
FRASE_PROLONGADOR_PLURAL = (
    "Os cabos prolongadores utilizados apresentam as seguintes "
    "resistências elétricas:"
)


def _prolongadores_da_inspecao(pacote: Pacote, padrao: float) -> list[float]:
    """Valores distintos de prolongador da inspeção (ordenados).

    Normalmente é um único valor (o mesmo cabo em todas as medições).
    """
    # Um zero vindo de máquina fora de faixa (">2000") não representa um cabo
    # e por isso fica de fora da lista.
    valores = sorted({
        e.prolongador for e in pacote.equipamentos
        if e.prolongador is not None
        and not (e.fora_de_faixa and e.prolongador == 0)
    })
    return valores or [padrao]


def _linha_prolongador_geral(doc, texto_valor: str, plural: bool = False) -> None:
    """Preenche a resistência do prolongador no laudo geral.

    O modelo traz a frase que a anuncia, mas não a linha com o valor; ela é
    inserida logo abaixo, com o mesmo estilo usado no laudo individual
    (negrito, 14 pt, justificado). Quando a inspeção tem mais de um valor de
    prolongador, a frase é ajustada para o plural.
    """
    from copy import deepcopy

    from docx.shared import Pt
    from docx.text.paragraph import Paragraph

    texto = f"Resistência elétrica do cabo prolongador = {texto_valor}"
    paragrafos = doc.paragraphs

    # Se a linha já existir no modelo, apenas atualiza o texto.
    for p in paragrafos:
        if p.text.startswith("Resistência elétrica do cabo prolongador ="):
            _substituir_no_paragrafo(p, p.text, texto)
            return

    marcador = "apresenta a seguinte resistência elétrica"
    for p in paragrafos:
        if marcador in p.text:
            if plural:
                _substituir_no_paragrafo(p, p.text, FRASE_PROLONGADOR_PLURAL)
            elemento = deepcopy(p._p)
            p._p.addnext(elemento)
            novo = Paragraph(elemento, p._parent)
            for run in list(novo.runs)[1:]:
                run._r.getparent().remove(run._r)
            if not novo.runs:
                novo.add_run("")
            run = novo.runs[0]
            run.text = texto
            run.bold = True
            run.font.size = Pt(14)
            return


# --- Laudo Geral ----------------------------------------------------------

def gerarLaudoGeral(
    pacote: Pacote,
    config: Configuracao,
    data_medicoes,
    caminho_saida: str,
    modelo: str | None = None,
    nome_planilha: str | None = None,
) -> str:
    """Gera o laudo geral (.docx) preenchendo o modelo.

    ``data_medicoes`` é um ``datetime.date`` (ou date-like) da inspeção.
    ``nome_planilha`` é o nome do arquivo da planilha resumo anexa (citado no
    corpo do laudo).
    """
    doc = docx.Document(modelo or MODELO_GERAL)
    data_ext = data_por_extenso(data_medicoes)

    subs = {
        # Capa
        "UNIDADE: INCUBATÓRIO – AURORA CHAPECÓ - SC": f"UNIDADE: {config.unidade}",
        "Chapecó – SC, 24 de julho de 2026": f"{config.cidade}, {data_ext}",
        # Engenheiro
        "Aníbal Rosa Vargas": config.engenheiro,
        "CREA-SC – 069788-5": f"CREA-SC – {config.crea}",
        # Nº da proposta (cabeçalho: "Proposta - <nº>")
        "102PC26AUR": config.proposta or "102PC26AUR",
        # Certificado de calibração — data e validade
        "certificado de calibração pelo fabricante na data 30/03/2026":
            f"certificado de calibração pelo fabricante na data "
            f"{config.calibracao_data}, com validade até {config.calibracao_validade}",
    }
    if nome_planilha:
        subs["Planilha resumo de resultado das medições - "
             "Aurora Fábrica Ração – Guatambu.xlsx"] = nome_planilha
    _aplicar_substituicoes(doc, subs)
    # Resistência do prolongador (valor vindo do pacote).
    valores_prol = _prolongadores_da_inspecao(pacote, config.prolongador_padrao)
    ohm = _ohm_do_documento(doc)
    _linha_prolongador_geral(
        doc,
        " e ".join(f"{_fmt(v)}m{ohm}" for v in valores_prol),
        plural=len(valores_prol) > 1,
    )
    _aplicar_imagens_config(doc, config)  # logo, equipamento, selo

    os.makedirs(os.path.dirname(os.path.abspath(caminho_saida)), exist_ok=True)
    doc.save(caminho_saida)
    return caminho_saida


def _nome_arquivo_seguro(texto: str) -> str:
    return texto.replace("/", "-").replace("\\", "-").replace(":", "-").strip()


# --- Laudo Individual -----------------------------------------------------

LIMITE_ADEQUADO = 1000.0  # mΩ (efetiva ≤ 1000 -> ESTÁ / adequado)


# Tamanho fixo das fotos do laudo individual (Figuras 2 e 3).
FOTO_LARGURA_CM = 8.0
FOTO_ALTURA_CM = 10.0


def _preparar_foto(caminho: str, formato: str, proporcao: float,
                   max_lado: int = 1400, q: int = 85) -> bytes:
    """Prepara a foto para ser exibida num quadro de proporção fixa.

    A imagem é encaixada (com preenchimento branco) na ``proporcao``
    ``largura/altura`` desejada, de modo que ao fixar largura e altura no Word
    ela não apareça distorcida.
    """
    from PIL import Image

    im = Image.open(caminho).convert("RGB")
    w, h = im.size
    if max(w, h) > max_lado:
        f = max_lado / max(w, h)
        im = im.resize((max(1, int(w * f)), max(1, int(h * f))), Image.LANCZOS)
        w, h = im.size

    # Quadro na proporção desejada, cobrindo a imagem inteira.
    if w / h > proporcao:          # imagem mais larga que o quadro
        quadro = (w, max(1, int(round(w / proporcao))))
    else:                          # imagem mais alta que o quadro
        quadro = (max(1, int(round(h * proporcao))), h)
    canvas = Image.new("RGB", quadro, (255, 255, 255))
    canvas.paste(im, ((quadro[0] - w) // 2, (quadro[1] - h) // 2))

    buf = io.BytesIO()
    canvas.save(buf, "PNG" if formato == "png" else "JPEG", quality=q)
    return buf.getvalue()


def _trocar_imagem(doc, inline_shape, caminho_foto: str) -> None:
    """Substitui a imagem de um inline shape pela foto, fixando o tamanho em
    ``FOTO_LARGURA_CM`` x ``FOTO_ALTURA_CM`` (sem distorcer o conteúdo)."""
    from docx.shared import Cm

    rId = inline_shape._inline.graphic.graphicData.pic.blipFill.blip.embed
    parte = doc.part.related_parts[rId]
    formato = "png" if "png" in (parte.content_type or "") else "jpeg"
    proporcao = FOTO_LARGURA_CM / FOTO_ALTURA_CM
    parte._blob = _preparar_foto(caminho_foto, formato, proporcao)
    inline_shape.width = Cm(FOTO_LARGURA_CM)
    inline_shape.height = Cm(FOTO_ALTURA_CM)


def _substituir_parte_imagem(doc, nome_parte: str, caminho_novo: str) -> bool:
    """Substitui o conteúdo da parte de imagem ``/word/media/<nome_parte>`` —
    atualiza todas as referências (capa e cabeçalhos compartilham a parte)."""
    alvo = f"/word/media/{nome_parte}"
    for part in doc.part.package.iter_parts():
        if str(part.partname) == alvo:
            formato = "png" if nome_parte.lower().endswith(".png") else "jpeg"
            part._blob = encaixar(caminho_novo, part.blob, formato)
            return True
    return False


def _aplicar_imagens_config(doc, config: Configuracao) -> None:
    """(d)(e) Substitui logo do cliente (capa e cabeçalhos), imagem do
    equipamento e selo de calibração pelas imagens fornecidas na configuração."""
    if config.logo_cliente and os.path.exists(config.logo_cliente):
        # image1.png: capa/cabeçalho (individual) e capa (geral).
        _substituir_parte_imagem(doc, "image1.png", config.logo_cliente)
        # image9.png: logo do cabeçalho do laudo geral.
        _substituir_parte_imagem(doc, "image9.png", config.logo_cliente)
    if config.imagem_equipamento and os.path.exists(config.imagem_equipamento):
        _substituir_parte_imagem(doc, "image2.jpeg", config.imagem_equipamento)
    if config.imagem_selo_calibracao and os.path.exists(config.imagem_selo_calibracao):
        _substituir_parte_imagem(doc, "image3.png", config.imagem_selo_calibracao)


def _legenda_figura1(doc, config: Configuracao, subs: dict) -> None:
    """Monta a legenda da Figura 1 com o instrumento e o modelo informados.

    No modelo a legenda é "Figura 1 – Miliohmímetro digital – MILLIOHM-1"; o
    nome do instrumento aparece com inicial maiúscula e o modelo com hífen,
    por isso ela é remontada em vez de substituída palavra a palavra.
    """
    import re

    instrumento = config.instrumento
    instrumento_cap = instrumento[:1].upper() + instrumento[1:]
    for p in doc.paragraphs:
        m = re.match(r"(Figura\s*1\s*([–-])\s*)", p.text)
        if m:
            traco = m.group(2)
            subs[p.text] = (f"{m.group(1)}{instrumento_cap} digital "
                            f"{traco} {config.instrumento_modelo}")
            return


def gerarLaudoIndividual(
    equipamento: Equipamento,
    config: Configuracao,
    data_medicoes,
    valor_medido: float | str,
    caminho_saida: str,
    prolongador: float | None = None,
    modelo: str | None = None,
) -> str:
    """Gera o laudo individual (.docx) de uma máquina preenchendo o modelo.

    ``valor_medido`` em mΩ, ou o texto ">2000" quando a medição ficou acima da
    escala do miliohmímetro. ``prolongador`` em mΩ; se None, usa o do
    equipamento ou o padrão da configuração.
    """
    doc = docx.Document(modelo or MODELO_INDIVIDUAL)
    if prolongador is None:
        prolongador = (
            equipamento.prolongador
            if equipamento.prolongador is not None
            else config.prolongador_padrao
        )
    # Resistência do próprio cabo (usada na frase da metodologia).
    prolongador_cabo = prolongador

    fora_faixa = valor_medido == FORA_DE_FAIXA
    if fora_faixa:
        # Acima da escala: nada a descontar e resultado inadequado.
        prolongador = 0.0
        texto_valor = FORA_DE_FAIXA
        texto_efetiva = FORA_DE_FAIXA
        adequado = False
    else:
        efetiva = valor_medido - prolongador
        texto_valor = _fmt(valor_medido)
        texto_efetiva = _fmt(efetiva)
        adequado = efetiva <= LIMITE_ADEQUADO
    data_ext = data_por_extenso(data_medicoes)
    # (a) Na capa, o nome do equipamento vai SEM o prefixo numérico "NN - ".
    nome_upper = equipamento.nome_sem_numero.upper()

    subs = {
        # Capa
        "UNIDADE: AURORA - INCUBATÓRIO BORMANN – CHAPECO - SC":
            f"UNIDADE: {config.unidade}",
        "INCUBADORA 2": nome_upper,
        "Chapecó – SC, 24 de julho de 2026": f"{config.cidade}, {data_ext}",
        # Engenheiro
        "Aníbal Rosa Vargas": config.engenheiro,
        "CREA-SC – 069788-5": f"CREA-SC – {config.crea}",
        # Nº da proposta (cabeçalho: "Proposta - <nº>")
        "102PC26AUR": config.proposta or "102PC26AUR",
        # Certificado de calibração — data e validade
        "certificado de calibração pelo fabricante na data 30/03/2026":
            f"certificado de calibração pelo fabricante na data "
            f"{config.calibracao_data}, com validade até {config.calibracao_validade}",
        # Instrumento de medição (só no laudo individual)
        "miliohmímetro": config.instrumento,
        "MILLIOHM 1": config.instrumento_modelo,
        "fabricante Instrument,": f"fabricante {config.instrumento_fabricante},",
        "correntes até 1,2 A": f"correntes até {config.instrumento_corrente}",
    }
    _legenda_figura1(doc, config, subs)
    # Linhas de MEDIÇÃO: resolve a string exata do modelo por prefixo (o modelo
    # usa o sinal de ohm U+2126, então reaproveitamos o caractere do próprio
    # texto em vez de digitá-lo).
    medicao = {
        "Valor medido = ": texto_valor,
        "Resistência elétrica do cabo prolongador PT = ": _fmt(prolongador),
        "Resistência elétrica de aterramento efetiva = ": texto_efetiva,
        # Frase da metodologia: resistência do próprio cabo prolongador.
        "Resistência elétrica do cabo prolongador = ": _fmt(prolongador_cabo),
    }
    for p in doc.paragraphs:
        for prefixo, valor in medicao.items():
            if p.text.startswith(prefixo):
                antigo = p.text
                ohm = antigo[-1]  # 'Ω' (U+2126) tal como no modelo
                subs[antigo] = f"{prefixo}{valor}m{ohm}"

    # (b)(c) Corrige as legendas das figuras (troca de descrição entre Fig. 2 e 3),
    # preservando o traço/espaçamento do próprio modelo.
    import re
    for p in doc.paragraphs:
        m2 = re.match(r"(Figura\s*2\s*[–-]\s*)", p.text)
        m3 = re.match(r"(Figura\s*3\s*[–-]\s*)", p.text)
        if m2:
            subs[p.text] = m2.group(1) + "Registro ensaio continuidade no equipamento"
        elif m3:
            subs[p.text] = m3.group(1) + "Resultado medição"
    if not adequado:
        subs["conclui-se que o equipamento  ESTÁ  solidamente conectado"] = (
            "conclui-se que o equipamento  NÃO ESTÁ  solidamente conectado"
        )
    _aplicar_substituicoes(doc, subs)

    # Troca as duas fotos variáveis (as duas últimas imagens do modelo):
    # penúltima = equipamento (foto 01), última = valor medido (foto 02).
    imagens = doc.inline_shapes
    foto_maquina = equipamento.fotos_maquina[0] if equipamento.fotos_maquina else None
    foto_valor = equipamento.fotos_valor[0] if equipamento.fotos_valor else None
    if len(imagens) >= 2:
        if foto_maquina and os.path.exists(foto_maquina):
            _trocar_imagem(doc, imagens[-2], foto_maquina)
        if foto_valor and os.path.exists(foto_valor):
            _trocar_imagem(doc, imagens[-1], foto_valor)

    _aplicar_imagens_config(doc, config)  # (d)(e) logo, equipamento, selo

    os.makedirs(os.path.dirname(os.path.abspath(caminho_saida)), exist_ok=True)
    doc.save(caminho_saida)
    return caminho_saida


def gerarLaudosIndividuais(
    pacote: Pacote,
    config: Configuracao,
    data_medicoes,
    medicoes: dict,
    pasta_saida: str,
) -> list[str]:
    """Gera um laudo individual (.docx) por máquina que tenha valor medido.

    ``medicoes``: ``{id_equipamento: {"valor": float, "prolongador": float|None}}``.
    Máquinas sem valor medido informado são puladas (não há como preencher a
    seção de medição). Devolve a lista de caminhos gerados.
    """
    os.makedirs(pasta_saida, exist_ok=True)
    gerados = []
    for eq in pacote.equipamentos:
        med = medicoes.get(eq.chave, {})
        valor = med.get("valor")
        if valor is None:
            continue
        # Nome do arquivo = nome da máquina (sem o prefixo "Laudo - ").
        nome = _nome_arquivo_seguro(eq.nome) or f"equipamento-{eq.chave}"
        caminho = os.path.join(pasta_saida, f"{nome}.docx")
        gerarLaudoIndividual(
            eq, config, data_medicoes, valor, caminho,
            prolongador=med.get("prolongador"),
        )
        gerados.append(caminho)
    return gerados

