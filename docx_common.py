#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
docx_common.py — utilidades compartilhadas entre reduzir_docx.py (1ª passada:
recompressão de fotos) e converter_png_jpeg.py (2ª passada: conversão de PNGs
fotográficos para JPEG).

Não é executado diretamente.
"""

import gc
import os
import zipfile

from PIL import Image
import docx  # python-docx

# Bomba de descompressão / sanidade: não tenta abrir imagens absurdas
Image.MAX_IMAGE_PIXELS = 300_000_000

DIM_MAXIMA = 1600            # maior dimensão (px) após redimensionamento
QUALIDADE_JPEG = 80          # qualidade de recompressão JPEG
TAMANHO_MINIMO = 200 * 1024  # imagens menores que isso são ignoradas (200 KB)
LIMITE_ALERTA = 100 * 1024 * 1024  # 100 MB — limiar do resumo final
CHUNK = 4 * 1024 * 1024      # cópia em blocos de 4 MB (economia de memória)

EXT_JPEG = {".jpg", ".jpeg", ".jpe", ".jfif"}
EXT_PNG = {".png"}
EXT_VETORIAL = {".emf", ".wmf", ".svg", ".wdp"}


def mb(n_bytes):
    return n_bytes / (1024 * 1024)


def fmt_mb(n_bytes):
    return f"{mb(n_bytes):,.1f} MB".replace(",", "X").replace(".", ",").replace("X", ".")


def eh_media(nome):
    return nome.replace("\\", "/").lower().startswith("word/media/")


def ext(nome):
    return os.path.splitext(nome)[1].lower()


def copiar_zipinfo(item):
    """Novo ZipInfo com os mesmos metadados de 'item', forçando deflate."""
    novo = zipfile.ZipInfo(filename=item.filename, date_time=item.date_time)
    novo.compress_type = zipfile.ZIP_DEFLATED
    novo.external_attr = item.external_attr
    novo.create_system = item.create_system
    return novo


def redimensionar_se_preciso(img, dim_maxima=DIM_MAXIMA):
    """Redimensiona 'img' para no máx. dim_maxima px na maior dimensão,
    mantendo a proporção. Retorna (imagem, foi_redimensionada)."""
    largura, altura = img.size
    maior = max(largura, altura)
    if maior <= dim_maxima:
        return img, False
    escala = dim_maxima / maior
    novo_tam = (max(1, round(largura * escala)), max(1, round(altura * escala)))
    if img.mode == "P":
        # Paleta: converte antes de redimensionar para não degradar cores
        img = img.convert("RGBA" if "transparency" in img.info else "RGB")
    return img.resize(novo_tam, Image.LANCZOS), True


def contagens_docx(caminho):
    """(nº parágrafos, nº tabelas, nº imagens inline) — usado na validação."""
    documento = docx.Document(caminho)
    contagens = (
        len(documento.paragraphs),
        len(documento.tables),
        len(documento.inline_shapes),
    )
    del documento
    gc.collect()
    return contagens
