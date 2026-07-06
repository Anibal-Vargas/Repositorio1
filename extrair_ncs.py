#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
extrair_ncs.py — Extrai TODAS as não conformidades (NCs) dos relatórios de
inspeção elétrica anonimizados (.docx) para uma única planilha Excel, SEM
resumir nem reescrever nada: o texto de cada campo é copiado integralmente.

Colunas da planilha gerada:
    arquivo_origem   — nome do .docx de onde a NC foi extraída
    secao            — último título/seção identificado antes da NC
                       (alta/média/baixa tensão, subestação, quadro etc.)
    descricao_nc     — texto integral da descrição da não conformidade
    itens_norma      — texto integral do embasamento normativo
                       (NR-10, NBR 5410, NBR 14039 etc.)
    solucao_proposta — texto integral da solução/recomendação proposta

Formatos de relatório reconhecidos (detecção automática, por arquivo):
  A) Tabela-matriz: cabeçalho com colunas do tipo
     "Item | Descrição da não conformidade | Embasamento | Solução proposta"
     e uma NC por linha;
  B) Blocos rotulados (em tabelas rótulo/valor ou em parágrafos corridos):
     "Não conformidade nº 01", "Descrição: ...", "Norma: ...",
     "Solução proposta: ..." — os rótulos são reconhecidos sem diferenciar
     maiúsculas/minúsculas nem acentos.

Regras de segurança:
  * Os relatórios são abertos SOMENTE para leitura; nada é alterado neles.
  * Nenhum texto é resumido, reescrito ou "corrigido" — cópia literal.
  * Fotos/espaços para foto são ignorados (não viram texto na planilha).
  * Ao final, o console e o log_extracao_ncs.txt informam: total de NCs,
    total por arquivo, arquivos com ZERO NCs (possível falha de extração)
    e NCs com campos vazios, para conferência manual.

Uso (Windows):
    py extrair_ncs.py
    py extrair_ncs.py --entrada "C:\\Temp\\RNCs\\Anonimizados" --saida "C:\\Temp\\RNCs\\Planilhas\\Extracao_NCs_bruta.xlsx"

Requisitos: Python 3.8+, python-docx, openpyxl
    py -m pip install -r requirements.txt
"""

import argparse
import os
import re
import sys
import traceback
import unicodedata
from datetime import datetime

try:
    import docx
    from docx.document import Document as _Documento
    from docx.oxml.ns import qn
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph
except ImportError:
    sys.exit("ERRO: python-docx não instalado. Execute:  py -m pip install python-docx")
try:
    import openpyxl
    from openpyxl.styles import Alignment, Font
    from openpyxl.utils import get_column_letter
except ImportError:
    sys.exit("ERRO: openpyxl não instalado. Execute:  py -m pip install openpyxl")

# ----------------------------- Configuração ---------------------------------

PASTA_ENTRADA_PADRAO = r"C:\Temp\RNCs\Anonimizados"
ARQUIVO_SAIDA_PADRAO = r"C:\Temp\RNCs\Planilhas\Extracao_NCs_bruta.xlsx"

COLUNAS = ["arquivo_origem", "secao", "descricao_nc", "itens_norma", "solucao_proposta"]

# Rótulos reconhecidos (comparados sem acentos/maiúsculas). A ordem dentro de
# cada categoria importa: os mais específicos vêm primeiro, para que
# "descrição da não conformidade" não seja capturado apenas como "descrição".
ROTULOS = {
    "descricao": [
        "descricao da nao conformidade", "descricao da irregularidade",
        "descricao do problema", "descricao da nc", "descricao",
        "nao conformidade encontrada", "nao conformidade constatada",
        "nao conformidade", "constatacao", "anomalia constatada",
        "anomalia", "irregularidade constatada", "irregularidade",
    ],
    "norma": [
        "embasamento tecnico normativo", "embasamento tecnico",
        "embasamento normativo", "embasamento", "fundamentacao normativa",
        "fundamentacao tecnica", "fundamentacao", "base normativa",
        "referencia normativa", "referencias normativas",
        "itens da norma", "item da norma", "itens de norma",
        "norma tecnica que a embasa", "norma tecnica", "normas tecnicas",
        "normas aplicaveis", "norma aplicavel", "normas", "norma",
    ],
    "solucao": [
        "solucao proposta", "solucoes propostas", "proposta de solucao",
        "proposta de correcao", "solucao recomendada", "solucao",
        "acao corretiva proposta", "acoes corretivas", "acao corretiva",
        "acao proposta", "medidas corretivas", "medida corretiva",
        "recomendacoes", "recomendacao", "providencias", "providencia",
        "correcao proposta", "adequacao proposta",
    ],
    "foto": [
        "registro fotografico", "registros fotograficos",
        "evidencia fotografica", "evidencias fotograficas",
        "fotografias", "fotografia", "fotos", "foto", "imagens", "imagem",
    ],
}

# Início de uma nova NC numerada, ex.: "NC 01", "N.C. 7", "Não Conformidade nº 12"
# (casa com o texto ORIGINAL — 'n[aã]o' cobre as formas com e sem acento)
RE_INICIO_NC = re.compile(
    r"^(nc|n\.c\.|n[aã]o\s+conformidade)\s*(n?[ºo°\.]?\s*)?(\d{1,4})\s*[-–—:.)]?\s*",
    re.IGNORECASE,
)

# Palavras-chave que caracterizam um título de seção do relatório
RE_SECAO = re.compile(
    r"\b(alta tensao|media tensao|baixa tensao|\bat\b|\bmt\b|\bbt\b|"
    r"subestacao|cabine primaria|cabine de medicao|cabine|transformador(es)?|"
    r"quadro(s)? (geral|gerais|de distribuicao|de forca|de comando|eletrico(s)?)|"
    r"qgbt|qdc|qdf|painel(eis)? eletrico(s)?|spda|aterramento|"
    r"iluminacao de emergencia|instalacoes eletricas)\b"
)

# ------------------------- Normalização de texto ----------------------------


def normalizar(texto):
    """Minúsculas, sem acentos e com espaços comprimidos — só para COMPARAR
    rótulos/títulos; o texto gravado na planilha é sempre o original."""
    texto = unicodedata.normalize("NFKD", texto)
    texto = "".join(c for c in texto if not unicodedata.combining(c))
    return re.sub(r"\s+", " ", texto).strip().lower()


def rotulo_prefixo(texto):
    """Se 'texto' COMEÇA com um rótulo conhecido, retorna (categoria, resto),
    onde 'resto' é o valor que segue o rótulo no texto ORIGINAL (após ':',
    '-' etc.). Caso contrário, retorna (None, None).

    Exige um separador (':', '-', '–', fim do texto ou quebra) logo após o
    rótulo, para não confundir com frases que apenas começam com a palavra
    (ex.: "Norma exige que..." não é rótulo)."""
    norm = normalizar(texto)
    for categoria, rotulos in ROTULOS.items():
        for rotulo in rotulos:
            if not norm.startswith(rotulo):
                continue
            resto_norm = norm[len(rotulo):]
            if resto_norm and not re.match(r"^\s*(n?[ºo°\.]?\s*\d{1,4})?\s*[:\-–—.)]", resto_norm):
                continue  # rótulo seguido de palavra comum: não é rótulo
            # Localiza no texto ORIGINAL o ponto equivalente ao fim do rótulo
            # + separador, para devolver o valor sem perder acentos/caixa.
            m = re.match(
                r"^\s*\S+(?:\s+\S+){%d}\s*(n?[ºo°\.]?\s*\d{1,4})?\s*[:\-–—.)]*\s*"
                % (len(rotulo.split()) - 1),
                texto,
            )
            resto = texto[m.end():].strip() if m else ""
            return categoria, resto
    return None, None


def rotulo_exato(texto):
    """Categoria cujo rótulo corresponde ao texto INTEIRO da célula (usado
    para reconhecer cabeçalhos de tabela-matriz); None se não corresponder."""
    norm = re.sub(r"[:\-–—.)]+$", "", normalizar(texto)).strip()
    norm = re.sub(r"\bn?[ºo°]\b$", "", norm).strip()
    for categoria, rotulos in ROTULOS.items():
        if norm in rotulos:
            return categoria
    return None


def eh_titulo_secao(paragrafo):
    """True se o parágrafo parece um título de seção (estilo Heading/Título
    ou texto curto com palavra-chave de seção elétrica)."""
    texto = paragrafo.text.strip()
    if not texto or len(texto) > 120:
        return False
    estilo = (paragrafo.style.name or "").lower() if paragrafo.style else ""
    if estilo.startswith(("heading", "título", "titulo", "title")):
        return True
    norm = normalizar(texto)
    if RE_SECAO.search(norm) and not rotulo_prefixo(texto)[0] and not RE_INICIO_NC.match(norm):
        # Título costuma ser curto e sem pontuação final de frase
        return len(norm) <= 90 and not norm.endswith((".", ";"))
    return False


# ------------------------ Percurso do documento ------------------------------


def iter_blocos(pai):
    """Percorre parágrafos e tabelas na ordem em que aparecem no documento
    (receita padrão do python-docx)."""
    if isinstance(pai, _Documento):
        elemento = pai.element.body
    elif isinstance(pai, _Cell):
        elemento = pai._tc
    else:
        raise TypeError("iter_blocos: tipo não suportado")
    for filho in elemento.iterchildren():
        if filho.tag == qn("w:p"):
            yield Paragraph(filho, pai)
        elif filho.tag == qn("w:tbl"):
            yield Table(filho, pai)


def texto_celula(celula):
    """Texto integral de uma célula (parágrafos unidos por quebra de linha)."""
    linhas = [p.text.strip() for p in celula.paragraphs]
    return "\n".join(l for l in linhas if l).strip()


def celulas_unicas(linha):
    """Células da linha sem repetições causadas por mesclagem."""
    vistos, unicas = set(), []
    for celula in linha.cells:
        if id(celula._tc) in vistos:
            continue
        vistos.add(id(celula._tc))
        unicas.append(celula)
    return unicas


# --------------------------- Montagem das NCs --------------------------------


class ColetorNCs:
    """Acumula os campos da NC em construção e a grava na lista ao 'fechar'.

    Mantém o estado necessário para os relatórios em blocos rotulados, em que
    descrição, norma e solução aparecem em parágrafos/linhas sucessivos."""

    def __init__(self, arquivo):
        self.arquivo = arquivo
        self.secao = ""
        self.ncs = []
        self._atual = None       # dict campos da NC em construção
        self._campo_ativo = None  # categoria que recebe parágrafos soltos

    def _nova(self):
        self._atual = {"descricao": [], "norma": [], "solucao": [], "secao": self.secao}
        self._desc_rotulada = False  # True quando a descrição veio de rótulo "Descrição:"

    def fechar(self):
        """Encerra a NC em construção, gravando-a se tiver algum conteúdo."""
        if self._atual and any(self._atual[c] for c in ("descricao", "norma", "solucao")):
            self.ncs.append({
                "arquivo_origem": self.arquivo,
                "secao": self._atual["secao"],
                "descricao_nc": "\n".join(self._atual["descricao"]).strip(),
                "itens_norma": "\n".join(self._atual["norma"]).strip(),
                "solucao_proposta": "\n".join(self._atual["solucao"]).strip(),
            })
        self._atual = None
        self._campo_ativo = None

    def definir_secao(self, texto):
        self.fechar()
        self.secao = texto.strip()

    def iniciar_nc(self, resto=""):
        self.fechar()
        self._nova()
        self._campo_ativo = "descricao"
        if resto:
            self._atual["descricao"].append(resto)

    def campo(self, categoria, valor):
        """Registra um rótulo encontrado (com valor possivelmente vazio)."""
        if categoria == "foto":
            self._campo_ativo = None  # ignora fotos/legendas que seguirem
            return
        # Um SEGUNDO rótulo "Descrição:" indica que começou OUTRA NC.
        # (Se a descrição atual veio só do título "NC nº X — ...", o rótulo
        # explícito pertence à MESMA NC e é acrescentado a ela.)
        if categoria == "descricao" and self._atual and self._desc_rotulada:
            self.fechar()
        if self._atual is None:
            self._nova()
        if categoria == "descricao":
            self._desc_rotulada = True
        self._campo_ativo = categoria
        if valor:
            self._atual[categoria].append(valor)

    def continuacao(self, texto):
        """Parágrafo sem rótulo: continua o campo ativo, se houver."""
        if self._atual is not None and self._campo_ativo:
            self._atual[self._campo_ativo].append(texto)

    def adicionar_pronta(self, descricao, norma, solucao, secao=None):
        """Grava uma NC completa (modo tabela-matriz), sem passar pelo estado."""
        self.ncs.append({
            "arquivo_origem": self.arquivo,
            "secao": (secao if secao is not None else self.secao).strip(),
            "descricao_nc": descricao.strip(),
            "itens_norma": norma.strip(),
            "solucao_proposta": solucao.strip(),
        })


# ------------------------- Extração das tabelas ------------------------------


def extrair_tabela_matriz(tabela, coletor):
    """Formato A: cabeçalho com ≥2 colunas rotuladas e uma NC por linha.
    Retorna True se a tabela foi reconhecida e processada nesse formato."""
    linhas = tabela.rows
    if len(linhas) < 2:
        return False
    cabecalho = linhas[0].cells
    mapa = {}  # índice da coluna -> categoria
    for i, celula in enumerate(cabecalho):
        categoria = rotulo_exato(texto_celula(celula))
        if categoria and categoria != "foto" and categoria not in mapa.values():
            mapa[i] = categoria
    categorias = set(mapa.values())
    if len(categorias & {"descricao", "norma", "solucao"}) < 2 or "descricao" not in categorias:
        return False
    for linha in linhas[1:]:
        celulas = linha.cells
        valores = {"descricao": "", "norma": "", "solucao": ""}
        for i, categoria in mapa.items():
            if i < len(celulas):
                valores[categoria] = texto_celula(celulas[i])
        if valores["descricao"]:
            coletor.adicionar_pronta(valores["descricao"], valores["norma"], valores["solucao"])
    return True


def extrair_tabela_rotulos(tabela, coletor):
    """Formato B em tabela: linhas/células no padrão rótulo -> valor
    ("Descrição: ..." na própria célula, ou rótulo numa célula e valor na
    seguinte). Retorna o nº de NCs fechadas durante o processamento."""
    antes = len(coletor.ncs)
    for linha in tabela.rows:
        celulas = celulas_unicas(linha)
        i = 0
        while i < len(celulas):
            celula = celulas[i]
            # Tabela aninhada dentro da célula: processa recursivamente
            for tab_interna in celula.tables:
                processar_tabela(tab_interna, coletor)
            texto = texto_celula(celula)
            if not texto:
                i += 1
                continue
            primeira_linha, _, resto_bloco = texto.partition("\n")
            m = RE_INICIO_NC.match(primeira_linha)
            if m and len(primeira_linha) <= 90:
                coletor.iniciar_nc(primeira_linha[m.end():].strip())
                if resto_bloco.strip():
                    coletor.continuacao(resto_bloco.strip())
                i += 1
                continue
            categoria, valor = rotulo_prefixo(primeira_linha)
            if categoria:
                if not valor and resto_bloco.strip():
                    valor = resto_bloco.strip()      # valor nas linhas de baixo da célula
                elif valor and resto_bloco.strip():
                    valor = valor + "\n" + resto_bloco.strip()
                if not valor and i + 1 < len(celulas):
                    valor = texto_celula(celulas[i + 1])  # valor na célula ao lado
                    i += 1
                coletor.campo(categoria, valor)
            else:
                coletor.continuacao(texto)
            i += 1
    return len(coletor.ncs) - antes


def processar_tabela(tabela, coletor):
    """Aplica primeiro o formato matriz; senão, o formato rótulo/valor."""
    try:
        if extrair_tabela_matriz(tabela, coletor):
            return
        extrair_tabela_rotulos(tabela, coletor)
    except Exception:
        # Tabela malformada não pode derrubar o arquivo inteiro
        pass


# --------------------------- Extração por arquivo ----------------------------


def extrair_arquivo(caminho):
    """Extrai as NCs de um .docx. Retorna (lista de NCs, None) ou ([], erro)."""
    nome = os.path.basename(caminho)
    try:
        documento = docx.Document(caminho)
    except Exception as e:
        return [], f"não foi possível abrir: {e}"
    coletor = ColetorNCs(nome)
    try:
        for bloco in iter_blocos(documento):
            if isinstance(bloco, Table):
                processar_tabela(bloco, coletor)
                continue
            texto = bloco.text.strip()
            if not texto:
                continue
            if eh_titulo_secao(bloco):
                coletor.definir_secao(texto)
                continue
            m = RE_INICIO_NC.match(texto)
            if m and len(texto) <= 200:
                coletor.iniciar_nc(texto[m.end():].strip())
                continue
            categoria, valor = rotulo_prefixo(texto)
            if categoria:
                coletor.campo(categoria, valor)
            else:
                coletor.continuacao(texto)
        coletor.fechar()
    except Exception:
        coletor.fechar()
        return coletor.ncs, f"erro durante a extração: {traceback.format_exc(limit=1).strip()}"
    return coletor.ncs, None


# ------------------------------- Planilha ------------------------------------


def gravar_planilha(ncs, caminho_saida):
    pasta = os.path.dirname(caminho_saida)
    if pasta:
        os.makedirs(pasta, exist_ok=True)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "NCs"
    ws.append(COLUNAS)
    for celula in ws[1]:
        celula.font = Font(bold=True)
    for nc in ncs:
        ws.append([nc[c] for c in COLUNAS])
    # Larguras confortáveis + quebra de linha nos campos de texto longo
    larguras = {"A": 40, "B": 30, "C": 70, "D": 50, "E": 60}
    for coluna, largura in larguras.items():
        ws.column_dimensions[coluna].width = largura
    quebra = Alignment(wrap_text=True, vertical="top")
    for linha in ws.iter_rows(min_row=2):
        for celula in linha:
            celula.alignment = quebra
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(COLUNAS))}{ws.max_row}"
    wb.save(caminho_saida)


# --------------------------------- Main --------------------------------------


def main():
    parser = argparse.ArgumentParser(
        description="Extrai as NCs de relatórios .docx para uma planilha única (cópia literal).")
    parser.add_argument("--entrada", default=PASTA_ENTRADA_PADRAO,
                        help=f"pasta com os .docx anonimizados (padrão: {PASTA_ENTRADA_PADRAO})")
    parser.add_argument("--saida", default=ARQUIVO_SAIDA_PADRAO,
                        help=f"arquivo .xlsx de saída (padrão: {ARQUIVO_SAIDA_PADRAO})")
    args = parser.parse_args()

    if not os.path.isdir(args.entrada):
        sys.exit(f"ERRO: pasta de entrada não encontrada: {args.entrada}")

    todos = sorted(os.listdir(args.entrada))
    arquivos_docx = [a for a in todos
                     if a.lower().endswith(".docx") and not a.startswith("~$")]
    ignorados = [a for a in todos
                 if a not in arquivos_docx
                 and os.path.isfile(os.path.join(args.entrada, a))]
    if not arquivos_docx:
        sys.exit(f"ERRO: nenhum .docx encontrado em {args.entrada}")

    print(f"Arquivos .docx encontrados: {len(arquivos_docx)}\n")
    todas_ncs = []
    por_arquivo = {}   # nome -> nº de NCs
    erros = {}         # nome -> mensagem de erro
    incompletas = {}   # nome -> nº de NCs com campo(s) vazio(s)

    for indice, nome in enumerate(arquivos_docx, 1):
        print(f"[{indice}/{len(arquivos_docx)}] {nome} ... ", end="", flush=True)
        ncs, erro = extrair_arquivo(os.path.join(args.entrada, nome))
        todas_ncs.extend(ncs)
        por_arquivo[nome] = len(ncs)
        if erro:
            erros[nome] = erro
        vazias = sum(1 for nc in ncs
                     if not nc["descricao_nc"] or not nc["itens_norma"]
                     or not nc["solucao_proposta"])
        if vazias:
            incompletas[nome] = vazias
        print(f"{len(ncs)} NC(s)" + (f"  [ERRO: {erro}]" if erro else ""))

    gravar_planilha(todas_ncs, args.saida)

    # ------------------------------ Resumo -----------------------------------
    sem_ncs = [n for n, q in por_arquivo.items() if q == 0]
    linhas = []
    linhas.append("=" * 78)
    linhas.append("RESUMO DA EXTRAÇÃO DE NCs — " + datetime.now().strftime("%d/%m/%Y %H:%M"))
    linhas.append("=" * 78)
    linhas.append(f"Pasta de entrada : {args.entrada}")
    linhas.append(f"Planilha gerada  : {args.saida}")
    linhas.append(f"Arquivos lidos   : {len(arquivos_docx)}")
    linhas.append(f"TOTAL DE NCs EXTRAÍDAS: {len(todas_ncs)}")
    linhas.append("")
    linhas.append("NCs por arquivo:")
    for nome in arquivos_docx:
        linhas.append(f"  {por_arquivo[nome]:>4}  {nome}")
    linhas.append("")
    linhas.append("CONFERIR MANUALMENTE (possível falha de extração):")
    if sem_ncs:
        linhas.append(f"  * {len(sem_ncs)} arquivo(s) com ZERO NCs encontradas:")
        for nome in sem_ncs:
            linhas.append(f"      - {nome}")
    if erros:
        linhas.append(f"  * {len(erros)} arquivo(s) com erro de leitura/extração:")
        for nome, msg in erros.items():
            linhas.append(f"      - {nome}: {msg}")
    if incompletas:
        linhas.append(f"  * {len(incompletas)} arquivo(s) com NCs de campo(s) vazio(s) "
                      "(descrição, norma ou solução em branco):")
        for nome, qtd in incompletas.items():
            linhas.append(f"      - {nome}: {qtd} NC(s) incompleta(s)")
    if ignorados:
        linhas.append(f"  * {len(ignorados)} arquivo(s) NÃO processado(s) por não serem .docx:")
        for nome in ignorados:
            linhas.append(f"      - {nome}")
    if not (sem_ncs or erros or incompletas or ignorados):
        linhas.append("  (nada a apontar — todos os arquivos geraram NCs completas)")
    resumo = "\n".join(linhas)
    print("\n" + resumo)

    caminho_log = os.path.join(os.path.dirname(args.saida) or ".", "log_extracao_ncs.txt")
    with open(caminho_log, "w", encoding="utf-8") as log:
        log.write(resumo + "\n")
    print(f"\nLog gravado em: {caminho_log}")


if __name__ == "__main__":
    main()
